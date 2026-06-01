# islam/scheduling/services/report_generator.py
"""Executive / Monthly Report Generator.

Generates professional PDF and DOCX reports from existing analytics services.
Every figure traces to a real service call — nothing is fabricated.

Supported formats
-----------------
  pdf   — HTML → PDF via pymupdf (fitz) Story renderer. No extra install needed.
  docx  — python-docx. Clean tables and styled headings.

Available sections (canonical order)
--------------------------------------
  executive_summary   Decision-layer prose (3 blocks)
  evm_performance     EVM KPI table + interpretation
  schedule_quality    DCMA 14-point check table
  delay_analysis      Top-10 delay root-causes
  risk_forecast       ML watchlist (near-critical, low P(on-time))
  anomalies           Counts + worst examples per category
  floor_health        Per-floor Build Quality vs Project Status
  cash_flow           Metrics + monthly disbursement table

Usage
-----
  data = gather_report_data(project_id, sections)
  pdf_bytes  = render_pdf(data, sections)
  docx_bytes = render_docx(data, sections)
"""

from __future__ import annotations

import io
import logging
from datetime import date
from typing import Any

logger = logging.getLogger(__name__)

# ── Section registry ──────────────────────────────────────────────────────────

ALL_SECTIONS: tuple[str, ...] = (
    "executive_summary",
    "evm_performance",
    "schedule_quality",
    "delay_analysis",
    "risk_forecast",
    "anomalies",
    "schedule_audit",
    "floor_health",
    "cash_flow",
)

SECTION_TITLES: dict[str, str] = {
    "executive_summary": "Executive Summary",
    "evm_performance": "EVM Performance",
    "schedule_quality": "Schedule Quality (DCMA 14-Point)",
    "delay_analysis": "Delay Root-Cause Analysis",
    "risk_forecast": "Risk Forecast (ML)",
    "anomalies": "Schedule Anomalies",
    "schedule_audit": "Schedule Audit — Section Mismatches",
    "floor_health": "Floor Health Matrix",
    "cash_flow": "Cash Flow Forecast",
}

REPORT_TYPES: dict[str, tuple[str, ...]] = {
    "executive": ALL_SECTIONS,
    # "weekly":  ("executive_summary", "evm_performance"),   # phase 2
    # "daily":   ("executive_summary",),                     # phase 2
}


# ── Helpers ───────────────────────────────────────────────────────────────────


def _fmt_num(v: float | None, decimals: int = 0) -> str:
    if v is None:
        return "n/a"
    if abs(v) >= 1_000_000:
        return f"{v / 1_000_000:,.2f}M"
    if abs(v) >= 1_000:
        return f"{v / 1_000:,.1f}k"
    return f"{v:,.{decimals}f}"


def _fmt_pct(v: float | None, decimals: int = 1) -> str:
    return "n/a" if v is None else f"{v:.{decimals}f}%"


def _fmt_date(iso: str | None) -> str:
    if not iso:
        return "n/a"
    try:
        d = date.fromisoformat(iso[:10])
        return d.strftime("%d %b %Y")
    except ValueError:
        return iso


def _forecast_finish(spi: float, project_start: str, project_end: str, as_of: str) -> date | None:
    """SPI-adjusted forecast finish (mirrors JS formula in evm.html)."""
    try:
        from datetime import timedelta

        spi = max(float(spi), 0.05)
        start = date.fromisoformat(project_start)
        baseline = date.fromisoformat(project_end)
        today = date.fromisoformat(as_of)
        elapsed = (today - start).days
        if today < baseline:
            remaining = (baseline - today).days
            fd = elapsed + remaining / spi
        else:
            fd = elapsed / spi
        return start + timedelta(days=max(0, round(fd)))
    except Exception:
        return None


def _safe(fn: Any, *args: Any, **kwargs: Any) -> dict:
    """Call a service function, returning {'has_data': False} on failure."""
    try:
        return fn(*args, **kwargs)  # type: ignore[operator]
    except Exception:
        logger.exception("report_generator: %s failed", getattr(fn, "__name__", fn))
        return {"has_data": False}


# ── Data gathering ────────────────────────────────────────────────────────────


def gather_report_data(project_id: str, sections: tuple[str, ...]) -> dict:
    """Call every service needed for the requested sections.

    Returns a single dict keyed by section name.  Missing / failed sections
    carry {"has_data": False} so the renderers can write 'Data not available'.
    """
    from environments.models import Project

    from .anomaly_detect import detect_anomalies
    from .cashflow import compute_cashflow
    from .completion_ml import run_completion_ml
    from .dcma_check import run_dcma_check
    from .decision_layer import compute_decision_summary
    from .delay_rootcause import run_delay_rootcause
    from .evm import compute_evm
    from .floor_health import compute_floor_health
    from .monte_carlo import compute_monte_carlo

    project = Project.objects.filter(pk=project_id).first()

    data: dict[str, Any] = {
        "_project_name": project.name if project else "Project",
        "_report_date": date.today().isoformat(),
        "_project_id": project_id,
    }

    # Always fetch EVM (needed for header)
    evm = _safe(compute_evm, project_id)
    data["evm"] = evm

    # Header fields derived from EVM
    data["_baseline_finish"] = evm.get("project_end")
    data["_as_of"] = evm.get("as_of")
    data["_forecast_finish"] = None
    if evm.get("has_data"):
        ff = _forecast_finish(evm["spi"], evm["project_start"], evm["project_end"], evm["as_of"])
        data["_forecast_finish"] = ff.isoformat() if ff else None

    if "executive_summary" in sections:
        data["decision"] = _safe(compute_decision_summary, project_id)

    if "schedule_quality" in sections:
        data["dcma"] = _safe(run_dcma_check, project_id)

    if "delay_analysis" in sections:
        data["drc"] = _safe(run_delay_rootcause, project_id)

    if "risk_forecast" in sections:
        data["ml"] = _safe(run_completion_ml, project_id)

    if "anomalies" in sections:
        data["anomaly"] = _safe(detect_anomalies, project_id)

    if "schedule_audit" in sections:
        from .schedule_audit import run_section_mismatch_audit

        # user=None → no LLM call; uses DB name cache from prior audit runs.
        # Degrades gracefully to {"has_data": False} when audit has never been run.
        data["audit"] = _safe(run_section_mismatch_audit, project_id, user=None)

    if "floor_health" in sections:
        data["floor"] = _safe(compute_floor_health, project_id)

    if "cash_flow" in sections:
        data["cf"] = _safe(compute_cashflow, project_id)

    # Monte Carlo (for executive summary baseline/P80 if available)
    if "executive_summary" in sections:
        data["mc"] = _safe(compute_monte_carlo, project_id)

    return data


# ── PDF renderer (pymupdf Story) ──────────────────────────────────────────────

_PDF_CSS = """
body {
    font-family: Helvetica, Arial, sans-serif;
    font-size: 9.5pt;
    color: #1a1a2e;
    margin: 0;
    padding: 0;
    line-height: 1.45;
}
h1  { font-size: 16pt; color: #1e3a5f; margin-bottom: 3pt; }
h2  { font-size: 11pt; color: #1e3a5f; margin-top: 0pt; margin-bottom: 4pt;
      border-bottom: 1pt solid #2563eb; padding-bottom: 2pt; }
p   { margin: 3pt 0; }
b   { font-weight: bold; }
i   { font-style: italic; }
table { width: 100%; border-collapse: collapse; margin-top: 5pt; }
th  { background-color: #1e3a5f; color: white; padding: 3.5pt 5pt;
      font-size: 8.5pt; text-align: left; font-weight: bold; }
td  { padding: 3pt 5pt; font-size: 8.5pt;
      border-bottom: 0.5pt solid #e2e8f0; vertical-align: top; }
.alt { background-color: #f8fafc; }
.good  { color: #16a34a; font-weight: bold; }
.bad   { color: #dc2626; font-weight: bold; }
.warn  { color: #d97706; font-weight: bold; }
.muted { color: #6b7280; }
.num   { text-align: right; font-variant-numeric: tabular-nums; }
.pb    { page-break-before: always; margin: 0; padding: 0; }
.label { font-size: 7.5pt; color: #6b7280; text-transform: uppercase;
         letter-spacing: 0.04em; margin-bottom: 1pt; }
.hdr-block { background-color: #f0f4ff; padding: 7pt 9pt;
             border-left: 3pt solid #2563eb; margin-bottom: 8pt; }
.na { color: #9ca3af; font-style: italic; }
"""


def _esc(text: str) -> str:
    """Minimal HTML escaping for report content."""
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _status_cls(status: str) -> str:
    return {"pass": "good", "warning": "warn", "fail": "bad"}.get(status, "muted")


def _sev_cls(spi_or_score: float, is_score: bool = False) -> str:
    if is_score:
        return "good" if spi_or_score >= 70 else "warn" if spi_or_score >= 40 else "bad"
    return "good" if spi_or_score >= 0.95 else "warn" if spi_or_score >= 0.80 else "bad"


def _html_header(data: dict) -> str:
    pn = _esc(data["_project_name"])
    rd = _fmt_date(data["_report_date"])
    aof = _fmt_date(data["_as_of"])
    bl = _fmt_date(data["_baseline_finish"])
    ff = _fmt_date(data["_forecast_finish"])
    evm = data.get("evm", {})
    spi = evm.get("spi")
    spi_str = f"SPI {spi:.2f}" if spi is not None else "SPI n/a"
    spi_cls = _sev_cls(spi) if spi is not None else "muted"
    return f"""
<h1>{pn} — Executive Report</h1>
<div class="hdr-block">
  <b>Report date:</b> {rd} &nbsp;&nbsp;
  <b>Data as of:</b> {aof} &nbsp;&nbsp;
  <b>Baseline finish:</b> {bl} &nbsp;&nbsp;
  <b>Forecast finish:</b> {ff} &nbsp;&nbsp;
  <span class="{spi_cls}">{spi_str}</span>
  <br/><span class="muted" style="font-size:7.5pt;">All values reflect original P6 coding.
  Audit-suggested reclassifications are not applied to reports.</span>
</div>
"""


def _html_executive_summary(data: dict) -> str:
    d = data.get("decision", {})
    if not d.get("has_data"):
        return "<p class='na'>Data not available.</p>"

    b1 = d["block1"]
    b2 = d["block2"]
    b3 = d["block3"]

    sev_cls = {"green": "good", "amber": "warn", "red": "bad"}

    def sev_label(sev: str) -> str:
        labels = {"green": "●", "amber": "▲", "red": "✕"}
        return f'<span class="{sev_cls.get(sev, "muted")}">{labels.get(sev, "")}</span>'

    html = ""
    # Block 1
    html += f"<p><b>{sev_label(b1.get('severity', ''))} {_esc(b1.get('sentence', ''))}</b></p>\n"
    # Block 2
    html += f"<p><b>{sev_label(b2.get('severity', ''))} {_esc(b2.get('sentence', ''))}</b></p>\n"

    # Block 3 — top-5 table
    items = b3.get("items", [])
    if items:
        html += "<p><b>Top 5 — Needs Attention</b></p>\n"
        html += "<table><tr><th>#</th><th>Activity</th><th>Type</th><th>Reason</th></tr>\n"
        for it in items:
            src = "Delay Driver" if it["source"] == "delay" else "At-Risk (ML)"
            html += (
                f"<tr><td class='num'>{it['rank']}</td>"
                f"<td>{_esc(it['name'])}</td>"
                f"<td>{_esc(src)}</td>"
                f"<td>{_esc(it['reason'])}</td></tr>\n"
            )
        html += "</table>\n"
    return html


def _html_evm(data: dict) -> str:
    evm = data.get("evm", {})
    if not evm.get("has_data"):
        return "<p class='na'>Data not available.</p>"

    use_cost = evm.get("use_cost", False)
    bac = evm.get("bac", 0)
    sym = "" if not use_cost else "$"

    def v(key: str, pct: bool = False) -> str:
        val = evm.get(key)
        if val is None:
            return "n/a"
        if pct:
            return _fmt_pct(val * 100 if abs(val) < 10 else val)
        return _fmt_num(val)

    spi = evm.get("spi", 1.0)
    cpi = evm.get("cpi", 1.0)
    rows = [
        ("BAC", f"{sym}{_fmt_num(bac)}", "Budget at Completion", "muted"),
        ("PV", f"{sym}{v('pv')}", "Planned Value (as-of)", "muted"),
        ("EV", f"{sym}{v('ev')}", "Earned Value", "muted"),
        ("AC", f"{sym}{v('ac')}", "Actual Cost", "muted"),
        ("SPI", f"{spi:.3f}", "Schedule Performance Index", _sev_cls(spi)),
        ("CPI", f"{cpi:.3f}", "Cost Performance Index", _sev_cls(cpi)),
        (
            "SV",
            f"{sym}{v('sv')}",
            "Schedule Variance",
            "good" if (evm.get("sv") or 0) >= 0 else "bad",
        ),
        ("CV", f"{sym}{v('cv')}", "Cost Variance", "good" if (evm.get("cv") or 0) >= 0 else "bad"),
        ("EAC", f"{sym}{v('eac')}", "Estimate at Completion", "muted"),
        (
            "VAC",
            f"{sym}{v('vac')}",
            "Variance at Completion",
            "good" if (evm.get("vac") or 0) >= 0 else "bad",
        ),
    ]
    html = "<table><tr><th>Metric</th><th>Value</th><th>Description</th></tr>\n"
    for i, (metric, val, desc, cls) in enumerate(rows):
        alt = ' class="alt"' if i % 2 else ""
        html += f"<tr{alt}><td><b>{_esc(metric)}</b></td><td class='num {cls}'>{_esc(val)}</td><td class='muted'>{_esc(desc)}</td></tr>\n"
    html += "</table>\n"

    # Interpretation sentence
    if spi >= 0.95 and cpi >= 0.95:
        interp = "Project is performing <b>on plan</b> — no corrective action indicated."
    elif spi < 0.80:
        months = 0.0
        ff = data.get("_forecast_finish")
        bl = data.get("_baseline_finish")
        if ff and bl:
            try:
                d_ff = date.fromisoformat(ff)
                d_bl = date.fromisoformat(bl)
                months = (d_ff - d_bl).days / 30.44
            except ValueError:
                pass
        mo_str = f", forecasting <b>{months:.0f} months late</b>" if months > 0.5 else ""
        interp = f"Schedule performance is <b class='bad'>significantly below plan</b>{mo_str}. Immediate recovery action required."
    else:
        interp = "Performance is <b class='warn'>below target</b> — monitor and apply corrective measures."

    html += f"<p><i>{interp}</i></p>\n"
    return html


def _html_schedule_quality(data: dict) -> str:
    dcma = data.get("dcma", {})
    if not dcma.get("has_data"):
        return "<p class='na'>Data not available.</p>"

    score = dcma.get("score", 0)
    color = dcma.get("score_color", "red")
    cls = _sev_cls(score, is_score=True)
    label = {"green": "Healthy", "amber": "Needs Attention", "red": "At Risk"}.get(color, color)
    s = dcma.get("summary", {})

    html = (
        f"<p><b>Overall score: <span class='{cls}'>{score}/100 — {label}</span></b> &nbsp;|&nbsp; "
    )
    html += f"<span class='good'>✓ {s.get('pass', 0)} Pass</span> &nbsp; "
    html += f"<span class='warn'>⚠ {s.get('warning', 0)} Warning</span> &nbsp; "
    html += f"<span class='bad'>✗ {s.get('fail', 0)} Fail</span></p>\n"

    html += "<table><tr><th>#</th><th>Check</th><th>Value</th><th>Target</th><th>Status</th></tr>\n"
    for i, chk in enumerate(dcma.get("checks", [])):
        alt = ' class="alt"' if i % 2 else ""
        sc = _status_cls(chk.get("status", ""))
        icon = {"pass": "✓", "warning": "⚠", "fail": "✗"}.get(chk.get("status", ""), "")
        val = chk.get("value", "n/a")
        if chk.get("unit") == "%":
            val = f"{val}%"
        html += (
            f"<tr{alt}><td class='num muted'>{i + 1}</td>"
            f"<td>{_esc(chk.get('name', ''))}"
            f"<br/><span class='muted' style='font-size:7.5pt;'>{_esc(chk.get('description', ''))}</span></td>"
            f"<td class='num'>{_esc(str(val))}</td>"
            f"<td class='muted'>{_esc(str(chk.get('target', '')))}</td>"
            f"<td class='{sc}'>{icon} {chk.get('status', '').upper()}</td></tr>\n"
        )
    html += "</table>\n"
    return html


def _html_delay_analysis(data: dict) -> str:
    drc = data.get("drc", {})
    if not drc.get("has_data"):
        return "<p class='na'>Data not available.</p>"

    s = drc.get("summary", {})
    rcs = drc.get("root_causes", [])

    html = (
        f"<p><b>{s.get('total_delayed', 0)}</b> delayed tasks &nbsp;|&nbsp; "
        f"<b>{s.get('root_causes', 0)}</b> root causes "
        f"({s.get('root_causes_firm', 0)} firm, "
        f"{s.get('root_causes_lower_confidence', 0)} lower-confidence) &nbsp;|&nbsp; "
        f"<b>{s.get('orphan', 0)}</b> orphan (untraceable) &nbsp;|&nbsp; "
        f"<b>{s.get('traceable_pct', 0)}%</b> traceable</p>\n"
    )

    html += "<table><tr><th>#</th><th>Root Cause Task</th><th>Type</th><th>Trade</th><th>Downstream</th><th>Chain-Days</th><th>Confidence</th></tr>\n"
    for i, rc in enumerate(rcs[:10]):
        alt = ' class="alt"' if i % 2 else ""
        conf = rc.get("confidence", "lower")
        cls = "good" if conf == "firm" else "warn"
        html += (
            f"<tr{alt}><td class='num muted'>{i + 1}</td>"
            f"<td>{_esc(rc.get('name', '')[:60])}</td>"
            f"<td class='muted'>{_esc(rc.get('activity_type', ''))}</td>"
            f"<td class='muted'>{_esc(rc.get('trade', ''))}</td>"
            f"<td class='num bad'>{rc.get('downstream_count', 0):,}</td>"
            f"<td class='num warn'>{rc.get('chain_delay_days', 0):,}</td>"
            f"<td class='{cls}'>{conf}</td></tr>\n"
        )
    html += "</table>\n"
    return html


def _html_risk_forecast(data: dict) -> str:
    ml = data.get("ml", {})
    if not ml.get("has_data"):
        return "<p class='na'>Data not available.</p>"

    mq = ml.get("model_quality", {})
    wl = ml.get("watchlist", [])
    auc = mq.get("auc", "n/a")
    auc_lbl = mq.get("auc_label", "")
    auc_cls = {"good": "good", "fair": "warn", "limited": "bad"}.get(auc_lbl, "muted")

    html = (
        f"<p><b>Model AUC: <span class='{auc_cls}'>{auc} ({auc_lbl})</span></b> &nbsp;|&nbsp; "
        f"Trained on {mq.get('n_train', 0):,} completed tasks &nbsp;|&nbsp; "
        f"Brier score {mq.get('brier', 'n/a')}</p>\n"
    )

    # Only near-critical tasks (float ≤ 10 wd)
    near_crit = [w for w in wl if w.get("near_critical", False)]
    if not near_crit:
        html += "<p class='na'>No near-critical incomplete tasks identified.</p>\n"
        return html

    html += "<table><tr><th>#</th><th>Task</th><th>Activity Code</th><th>Float (wd)</th><th>P(on-time)</th><th>Due</th></tr>\n"
    for i, w in enumerate(near_crit[:15]):
        alt = ' class="alt"' if i % 2 else ""
        prob = w.get("probability_on_time", 0.5)
        pct = round(prob * 100)
        pc = "good" if pct >= 70 else "warn" if pct >= 50 else "bad"
        tf = w.get("total_float", 9999)
        tf_s = str(tf) if tf < 9000 else "—"
        tc = "bad" if tf < 0 else "warn" if tf <= 5 else "muted"
        code = w.get("activity_code", "")
        html += (
            f"<tr{alt}><td class='num muted'>{i + 1}</td>"
            f"<td>{_esc(w.get('name', '')[:55])}</td>"
            f"<td class='muted' style='font-size:7.5pt;'>{_esc(code)}</td>"
            f"<td class='num {tc}'>{tf_s}</td>"
            f"<td class='num {pc}'>{pct}%</td>"
            f"<td class='muted'>{_fmt_date(w.get('baseline_finish', ''))}</td></tr>\n"
        )
    html += "</table>\n"
    return html


def _html_anomalies(data: dict) -> str:
    an = data.get("anomaly", {})
    if not an.get("has_data"):
        return "<p class='na'>Data not available.</p>"

    s = an.get("summary", {})
    by = s.get("by_type", {})
    html = (
        f"<p><b>{s.get('total_flagged', 0)}</b> tasks flagged "
        f"({s.get('flagged_pct', 0)}% of {s.get('total_tasks', 0):,}) &nbsp;|&nbsp; "
        f"Running long: <b>{by.get('running_long', 0)}</b> &nbsp; "
        f"Unrealistic baseline: <b>{by.get('unrealistic_baseline', 0)}</b> &nbsp; "
        f"Statistical outliers: <b>{by.get('statistical_outlier', 0)}</b> &nbsp; "
        f"Logic: <b>{by.get('logic_anomaly', 0)}</b></p>\n"
    )

    def _worst_table(items: list, label: str, max_n: int = 5) -> str:
        if not items:
            return ""
        t = f"<p><b>{_esc(label)} (top {min(len(items), max_n)})</b></p>\n"
        t += "<table><tr><th>Task</th><th>Trade</th><th>Detail</th></tr>\n"
        for i, it in enumerate(items[:max_n]):
            alt = ' class="alt"' if i % 2 else ""
            t += (
                f"<tr{alt}><td>{_esc(it.get('name', '')[:55])}</td>"
                f"<td class='muted'>{_esc(it.get('trade', ''))}</td>"
                f"<td class='warn' style='font-size:7.5pt;'>{_esc(it.get('explanation', ''))}</td></tr>\n"
            )
        t += "</table>\n"
        return t

    html += _worst_table(an.get("running_long", []), "Running Long (≥2× planned duration)")
    html += _worst_table(an.get("statistical_outliers", []), "Statistical Outliers")
    html += _worst_table(an.get("logic_anomalies", []), "Logic / Data Quality Errors")
    return html


def _html_floor_health(data: dict) -> str:
    fh = data.get("floor", {})
    if not fh.get("has_data"):
        return "<p class='na'>Data not available.</p>"

    floors = fh.get("floors", [])
    html = (
        "<p><i>Build Quality = schedule construction quality (logic gaps + structural anomalies). "
    )
    html += "Project Status = work progress (delay rate + ML forward risk). "
    html += "Scores 0-100; higher = better.</i></p>\n"

    def score_cls(s: int) -> str:
        return "good" if s >= 70 else "warn" if s >= 40 else "bad"

    html += "<table><tr><th>Floor</th><th>Tasks</th>"
    html += "<th>Build Quality</th><th>Defect Rate</th><th>Logic Issues</th><th>Anomaly Flags</th>"
    html += "<th>Project Status</th><th>Delayed %</th><th>Avg P(on-time)</th></tr>\n"
    for i, f in enumerate(floors):
        alt = ' class="alt"' if i % 2 else ""
        bq = f.get("build_quality", {})
        ps = f.get("project_status", {})
        bq_s = bq.get("score", 0)
        ps_s = ps.get("score", 0)
        prob = ps.get("avg_prob_on_time")
        prob_s = f"{round(prob * 100)}%" if prob is not None else "n/a"
        html += (
            f"<tr{alt}><td><b>{_esc(f.get('token', ''))}</b>"
            f"<br/><span class='muted' style='font-size:7pt;'>{_esc(f.get('label', ''))}</span></td>"
            f"<td class='num muted'>{f.get('task_count', 0)}</td>"
            f"<td class='num {score_cls(bq_s)}'>{bq_s}</td>"
            f"<td class='num muted'>{bq.get('defect_rate_pct', '—')}%</td>"
            f"<td class='num muted'>{bq.get('logic_issues', 0)}</td>"
            f"<td class='num muted'>{bq.get('anomaly_flags', 0)}</td>"
            f"<td class='num {score_cls(ps_s)}'>{ps_s}</td>"
            f"<td class='num muted'>{ps.get('delay_pct', '—')}%</td>"
            f"<td class='num muted'>{prob_s}</td></tr>\n"
        )
    html += "</table>\n"
    return html


def _html_cash_flow(data: dict) -> str:
    cf = data.get("cf", {})
    if not cf.get("has_data"):
        return "<p class='na'>Data not available.</p>"

    m = cf.get("metrics", {})
    src = cf.get("source", "")
    src_label = "P6 resource assignments" if src == "p6_assignments" else "task costs"
    months = cf.get("months", [])

    def fc(v: float | None) -> str:
        return _fmt_num(v) if v is not None else "—"

    html = (
        f"<p><b>BAC:</b> {fc(m.get('bac'))} &nbsp;|&nbsp; "
        f"<b>Spent:</b> {fc(m.get('ac'))} ({m.get('spent_pct', '—')}%) &nbsp;|&nbsp; "
        f"<b>Remaining:</b> {fc(m.get('remaining'))} &nbsp;|&nbsp; "
        f"<b>Peak month:</b> {m.get('peak_month', '—')} ({fc(m.get('peak_value'))}) &nbsp;|&nbsp; "
        f"Source: {src_label}</p>\n"
    )

    # Monthly table — cap at 36 months to keep report readable
    visible = months[:36]
    html += "<table><tr><th>Month</th><th>Planned</th><th>Actual</th><th>Forecast</th><th>Cum. Planned</th><th>Cum. Progress</th></tr>\n"
    for i, mo in enumerate(visible):
        alt = ' class="alt"' if i % 2 else ""
        html += (
            f"<tr{alt}><td class='muted'>{_esc(mo.get('month', ''))}</td>"
            f"<td class='num'>{fc(mo.get('planned'))}</td>"
            f"<td class='num'>{fc(mo.get('actual')) if mo.get('actual') else '—'}</td>"
            f"<td class='num muted'>{fc(mo.get('forecast')) if mo.get('forecast') else '—'}</td>"
            f"<td class='num'>{fc(mo.get('cumulative_planned'))}</td>"
            f"<td class='num'>{fc(mo.get('cumulative_progress'))}</td></tr>\n"
        )
    if len(months) > 36:
        html += "<tr><td colspan='6' class='muted' style='font-style:italic;padding:4pt;'>"
        html += f"… {len(months) - 36} additional months not shown (cap: 36 for print readability)</td></tr>\n"
    html += "</table>\n"
    return html


def _html_schedule_audit(data: dict) -> str:
    audit = data.get("audit", {})

    advisory = (
        "<p class='muted' style='border-left:3pt solid #9ca3af;padding-left:6pt;margin-bottom:6pt;'>"
        "<b>Advisory only.</b> Flagged items are for human review — the original P6 CSI coding "
        "is the source of record and has <b>not</b> been changed by this tool. "
        "Confirmed items are high-confidence disagreements between the activity name and its "
        "coded CSI division. Needs-review items have medium confidence and may reflect a "
        "legitimate planner decision.</p>\n"
    )

    if not audit.get("has_data"):
        return (
            advisory
            + "<p class='na'>Schedule audit data not available. Run the Schedule Audit in the "
            "EVM tab to generate findings; results are cached in the database for subsequent "
            "report runs.</p>\n"
        )

    items: list[dict] = audit.get("items", [])
    confirmed = [it for it in items if it.get("verdict") == "confirmed"]
    needs_review = [it for it in items if it.get("verdict") == "needs_review"]
    ai_ran = audit.get("ai_ran", False)

    html = advisory

    n_conf = audit.get("confirmed_count", len(confirmed))
    n_rev = audit.get("needs_review_count", len(needs_review))
    html += (
        f"<p><b>{n_conf} confirmed</b> coding mismatch{'' if n_conf == 1 else 'es'} · "
        f"<b>{n_rev} for review</b> · "
        f"{'AI classification active' if ai_ran else 'Classification from cached results'}</p>\n"
    )

    def _verdict_cls(v: str) -> str:
        return {"confirmed": "bad", "needs_review": "warn"}.get(v, "muted")

    def _conf_cls(c: str) -> str:
        return {"high": "good", "medium": "warn", "low": "muted"}.get(c or "", "muted")

    def _audit_table(rows: list[dict], label: str) -> str:
        if not rows:
            return ""
        t = f"<p><b>{_esc(label)}</b></p>\n"
        t += (
            "<table><tr>"
            "<th>Activity name</th>"
            "<th>Code</th>"
            "<th>Coded div.</th>"
            "<th>Suggested div.</th>"
            "<th>Conf.</th>"
            "<th>Reason</th>"
            "</tr>\n"
        )
        for i, it in enumerate(rows):
            alt = ' class="alt"' if i % 2 else ""
            vc = _verdict_cls(it.get("verdict", ""))
            cc = _conf_cls(it.get("ai_confidence", ""))
            t += (
                f"<tr{alt}>"
                f"<td>{_esc(it.get('name', '')[:55])}</td>"
                f"<td class='muted' style='font-size:7.5pt;'>{_esc(it.get('activity_code', ''))}</td>"
                f"<td class='{vc}'>{_esc(it.get('coded_trade', ''))}</td>"
                f"<td><b>{_esc(it.get('ai_trade', '') or '—')}</b></td>"
                f"<td class='{cc}'>{_esc(it.get('ai_confidence', '') or '—')}</td>"
                f"<td class='muted' style='font-size:7.5pt;'>{_esc(it.get('ai_reason', '') or '')}</td>"
                f"</tr>\n"
            )
        t += "</table>\n"
        return t

    html += _audit_table(confirmed, f"Confirmed mismatches ({len(confirmed)})")
    html += _audit_table(needs_review, f"Needs review ({len(needs_review)})")

    if not confirmed and not needs_review:
        html += "<p class='na'>No confirmed or needs-review items in the stored audit result.</p>\n"

    return html


_SECTION_HTML_BUILDERS: dict[str, Any] = {
    "executive_summary": _html_executive_summary,
    "evm_performance": _html_evm,
    "schedule_quality": _html_schedule_quality,
    "delay_analysis": _html_delay_analysis,
    "risk_forecast": _html_risk_forecast,
    "anomalies": _html_anomalies,
    "schedule_audit": _html_schedule_audit,
    "floor_health": _html_floor_health,
    "cash_flow": _html_cash_flow,
}


def _build_full_html(data: dict, sections: tuple[str, ...]) -> str:
    parts = [f"<style>{_PDF_CSS}</style>\n", _html_header(data)]
    sec_num = 1
    for i, key in enumerate(sections):
        if key not in _SECTION_HTML_BUILDERS:
            continue
        title = SECTION_TITLES.get(key, key)
        pb = '<p class="pb"></p>\n' if i > 0 else ""
        parts.append(f"{pb}<h2>{sec_num}. {_esc(title)}</h2>\n")
        parts.append(_SECTION_HTML_BUILDERS[key](data))
        sec_num += 1

    # Footer
    rd = _fmt_date(data["_report_date"])
    pn = _esc(data["_project_name"])
    parts.append(
        f"<p style='margin-top:18pt;font-size:7.5pt;color:#9ca3af;border-top:0.5pt solid #e5e7eb;"
        f"padding-top:4pt;'>Generated by Castor · {pn} · {rd} · All figures from live schedule data.</p>\n"
    )
    return "".join(parts)


def render_pdf(data: dict, sections: tuple[str, ...]) -> bytes:
    """Render the report to PDF bytes using pymupdf Story."""
    import fitz

    html = _build_full_html(data, sections)
    buf = io.BytesIO()
    story = fitz.Story(html)
    writer = fitz.DocumentWriter(buf)

    mediabox = fitz.paper_rect("A4")
    margin_x, margin_y = 55, 55
    where = mediabox + (margin_x, margin_y, -margin_x, -margin_y)

    more = 1
    while more:
        dev = writer.begin_page(mediabox)
        more, _ = story.place(where)
        story.draw(dev)
        writer.end_page()

    writer.close()
    return buf.getvalue()


# ── DOCX renderer (python-docx) ───────────────────────────────────────────────


def _docx_heading(doc: Any, text: str, level: int = 1) -> None:
    """Add a heading paragraph to the document."""
    from docx.shared import Pt, RGBColor

    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    run.font.bold = True
    run.font.size = Pt(14 if level == 1 else 11)


def _docx_table(doc: Any, headers: list[str], rows: list[list[str]]) -> None:
    """Add a styled table to the document."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"

    # Header row
    hdr_row = tbl.rows[0]
    for j, h in enumerate(headers):
        cell = hdr_row.cells[j]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Blue background via XML
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1E3A5F")
        tc_pr.append(shd)

    # Data rows
    for i, row in enumerate(rows):
        tr = tbl.add_row()
        for j, cell_text in enumerate(row):
            cell = tr.cells[j]
            cell.text = cell_text
            p = cell.paragraphs[0]
            if p.runs:
                p.runs[0].font.size = Pt(8.5)
            if i % 2 == 1:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F8FAFC")
                tc_pr.append(shd)


def render_docx(data: dict, sections: tuple[str, ...]) -> bytes:
    """Render the report to DOCX bytes using python-docx."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()

    # Remove the bare <w:zoom> from the default template's settings.xml.
    # python-docx emits <w:zoom w:val="bestFit"/> without w:percent, which
    # fails strict OOXML schema validation (Word tolerates it, validators don't).
    from docx.oxml.ns import qn as _qn

    _settings_elem = doc.settings.element
    for _z in _settings_elem.findall(_qn("w:zoom")):
        _settings_elem.remove(_z)

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Header ────────────────────────────────────────────────────────────────
    pn = data["_project_name"]
    rd = _fmt_date(data["_report_date"])
    aof = _fmt_date(data["_as_of"])
    bl = _fmt_date(data["_baseline_finish"])
    ff = _fmt_date(data["_forecast_finish"])
    evm = data.get("evm", {})
    spi = evm.get("spi")

    title = doc.add_paragraph()
    run = title.add_run(f"{pn} — Executive Report")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    meta = doc.add_paragraph()
    meta.add_run(
        f"Report date: {rd}  |  Data as of: {aof}  |  Baseline finish: {bl}  |  Forecast finish: {ff}"
    )
    if spi is not None:
        meta.add_run(f"  |  SPI: {spi:.2f}")
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    note = doc.add_paragraph(
        "All values reflect original P6 coding. "
        "Audit-suggested reclassifications are not applied to reports."
    )
    note.runs[0].font.size = Pt(7.5)
    note.runs[0].font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    doc.add_paragraph()

    sec_num = 1

    for i, key in enumerate(sections):
        if key not in _SECTION_HTML_BUILDERS:
            continue

        title_text = f"{sec_num}. {SECTION_TITLES.get(key, key)}"
        sec_num += 1

        if i > 0:
            doc.add_page_break()

        _docx_heading(doc, title_text, level=1)

        # ── Section content ────────────────────────────────────────────────
        if key == "executive_summary":
            d = data.get("decision", {})
            if not d.get("has_data"):
                doc.add_paragraph("Data not available.")
            else:
                b1, b2, b3 = d["block1"], d["block2"], d["block3"]
                for blk in (b1, b2):
                    p = doc.add_paragraph()
                    p.add_run(blk.get("sentence", "")).font.size = Pt(9.5)
                items = b3.get("items", [])
                if items:
                    doc.add_paragraph("Top 5 — Needs Attention").runs[0].bold = True
                    _docx_table(
                        doc,
                        ["#", "Activity", "Type", "Reason"],
                        [
                            [
                                str(it["rank"]),
                                it["name"][:60],
                                "Delay Driver" if it["source"] == "delay" else "At-Risk (ML)",
                                it["reason"],
                            ]
                            for it in items
                        ],
                    )

        elif key == "evm_performance":
            e = data.get("evm", {})
            if not e.get("has_data"):
                doc.add_paragraph("Data not available.")
            else:
                use_cost = e.get("use_cost", False)
                sym = "" if not use_cost else "$"
                _docx_table(
                    doc,
                    ["Metric", "Value", "Description"],
                    [
                        ["BAC", f"{sym}{_fmt_num(e.get('bac', 0))}", "Budget at Completion"],
                        ["PV", f"{sym}{_fmt_num(e.get('pv', 0))}", "Planned Value"],
                        ["EV", f"{sym}{_fmt_num(e.get('ev', 0))}", "Earned Value"],
                        ["AC", f"{sym}{_fmt_num(e.get('ac', 0))}", "Actual Cost"],
                        ["SPI", f"{e.get('spi', 0):.3f}", "Schedule Performance Index"],
                        ["CPI", f"{e.get('cpi', 0):.3f}", "Cost Performance Index"],
                        ["SV", f"{sym}{_fmt_num(e.get('sv', 0))}", "Schedule Variance"],
                        ["CV", f"{sym}{_fmt_num(e.get('cv', 0))}", "Cost Variance"],
                        ["EAC", f"{sym}{_fmt_num(e.get('eac', 0))}", "Estimate at Completion"],
                        ["VAC", f"{sym}{_fmt_num(e.get('vac', 0))}", "Variance at Completion"],
                    ],
                )

        elif key == "schedule_quality":
            dcma = data.get("dcma", {})
            if not dcma.get("has_data"):
                doc.add_paragraph("Data not available.")
            else:
                s = dcma.get("summary", {})
                p = doc.add_paragraph()
                p.add_run(
                    f"Score: {dcma.get('score', 0)}/100  |  "
                    f"Pass: {s.get('pass', 0)}  Warning: {s.get('warning', 0)}  Fail: {s.get('fail', 0)}"
                )
                _docx_table(
                    doc,
                    ["#", "Check", "Value", "Target", "Status"],
                    [
                        [
                            str(i + 1),
                            chk.get("name", ""),
                            str(chk.get("value", "")),
                            str(chk.get("target", "")),
                            chk.get("status", "").upper(),
                        ]
                        for i, chk in enumerate(dcma.get("checks", []))
                    ],
                )

        elif key == "delay_analysis":
            drc = data.get("drc", {})
            if not drc.get("has_data"):
                doc.add_paragraph("Data not available.")
            else:
                s = drc.get("summary", {})
                doc.add_paragraph(
                    f"{s.get('total_delayed', 0)} delayed  |  "
                    f"{s.get('root_causes', 0)} root causes  |  "
                    f"{s.get('orphan', 0)} orphan  |  {s.get('traceable_pct', 0)}% traceable"
                )
                _docx_table(
                    doc,
                    [
                        "#",
                        "Root Cause Task",
                        "Type",
                        "Trade",
                        "Downstream",
                        "Chain-Days",
                        "Confidence",
                    ],
                    [
                        [
                            str(i + 1),
                            rc.get("name", "")[:55],
                            rc.get("activity_type", ""),
                            rc.get("trade", ""),
                            str(rc.get("downstream_count", 0)),
                            str(rc.get("chain_delay_days", 0)),
                            rc.get("confidence", ""),
                        ]
                        for i, rc in enumerate(drc.get("root_causes", [])[:10])
                    ],
                )

        elif key == "risk_forecast":
            ml = data.get("ml", {})
            if not ml.get("has_data"):
                doc.add_paragraph("Data not available.")
            else:
                mq = ml.get("model_quality", {})
                wl = [w for w in ml.get("watchlist", []) if w.get("near_critical", False)]
                doc.add_paragraph(
                    f"Model AUC: {mq.get('auc', 'n/a')} ({mq.get('auc_label', '')})  |  "
                    f"Brier: {mq.get('brier', 'n/a')}  |  Trained on {mq.get('n_train', 0):,} tasks"
                )
                if wl:
                    _docx_table(
                        doc,
                        ["#", "Task", "Activity Code", "Float (wd)", "P(on-time)", "Due"],
                        [
                            [
                                str(i + 1),
                                w.get("name", "")[:55],
                                w.get("activity_code", ""),
                                str(w.get("total_float", "—")),
                                f"{round(w.get('probability_on_time', 0) * 100)}%",
                                _fmt_date(w.get("baseline_finish", "")),
                            ]
                            for i, w in enumerate(wl[:15])
                        ],
                    )
                else:
                    doc.add_paragraph("No near-critical incomplete tasks identified.")

        elif key == "anomalies":
            an = data.get("anomaly", {})
            if not an.get("has_data"):
                doc.add_paragraph("Data not available.")
            else:
                s = an.get("summary", {})
                by = s.get("by_type", {})
                doc.add_paragraph(
                    f"{s.get('total_flagged', 0)} flagged ({s.get('flagged_pct', 0)}%)  |  "
                    f"Running long: {by.get('running_long', 0)}  "
                    f"Unrealistic baseline: {by.get('unrealistic_baseline', 0)}  "
                    f"Outliers: {by.get('statistical_outlier', 0)}  "
                    f"Logic: {by.get('logic_anomaly', 0)}"
                )
                for label, items in [
                    ("Running Long", an.get("running_long", [])[:5]),
                    ("Logic / Data Quality", an.get("logic_anomalies", [])[:5]),
                ]:
                    if items:
                        doc.add_paragraph(label).runs[0].bold = True
                        _docx_table(
                            doc,
                            ["Task", "Trade", "Detail"],
                            [
                                [
                                    it.get("name", "")[:55],
                                    it.get("trade", ""),
                                    it.get("explanation", "")[:80],
                                ]
                                for it in items
                            ],
                        )

        elif key == "schedule_audit":
            au = data.get("audit", {})
            advisory_text = (
                "Advisory only. Flagged items are for human review — original P6 CSI coding "
                "is the source of record and has not been changed."
            )
            p = doc.add_paragraph(advisory_text)
            p.runs[0].italic = True
            p.runs[0].font.size = Pt(8.5)
            if not au.get("has_data"):
                doc.add_paragraph(
                    "Schedule audit data not available. Run the Schedule Audit in the EVM tab "
                    "to generate findings."
                )
            else:
                items: list[dict] = au.get("items", [])
                confirmed_items = [it for it in items if it.get("verdict") == "confirmed"]
                review_items = [it for it in items if it.get("verdict") == "needs_review"]
                doc.add_paragraph(
                    f"{au.get('confirmed_count', len(confirmed_items))} confirmed  |  "
                    f"{au.get('needs_review_count', len(review_items))} needs review"
                )
                for label, group in [
                    ("Confirmed Mismatches", confirmed_items),
                    ("Needs Review", review_items),
                ]:
                    if not group:
                        continue
                    doc.add_paragraph(label).runs[0].bold = True
                    _docx_table(
                        doc,
                        [
                            "Activity Name",
                            "Code",
                            "Coded Div.",
                            "Suggested Div.",
                            "Conf.",
                            "Reason",
                        ],
                        [
                            [
                                it.get("name", "")[:55],
                                it.get("activity_code", ""),
                                it.get("coded_trade", ""),
                                it.get("ai_trade", "") or "—",
                                it.get("ai_confidence", "") or "—",
                                (it.get("ai_reason", "") or "")[:80],
                            ]
                            for it in group
                        ],
                    )

        elif key == "floor_health":
            fh = data.get("floor", {})
            if not fh.get("has_data"):
                doc.add_paragraph("Data not available.")
            else:
                floors = fh.get("floors", [])
                _docx_table(
                    doc,
                    [
                        "Floor",
                        "Tasks",
                        "Build Quality",
                        "Defect Rate",
                        "Logic Issues",
                        "Anomaly Flags",
                        "Proj. Status",
                        "Delayed %",
                        "Avg P(on-time)",
                    ],
                    [
                        [
                            f.get("token", ""),
                            str(f.get("task_count", 0)),
                            str(f["build_quality"].get("score", 0)),
                            f"{f['build_quality'].get('defect_rate_pct', '—')}%",
                            str(f["build_quality"].get("logic_issues", 0)),
                            str(f["build_quality"].get("anomaly_flags", 0)),
                            str(f["project_status"].get("score", 0)),
                            f"{f['project_status'].get('delay_pct', '—')}%",
                            (
                                f"{round(f['project_status']['avg_prob_on_time'] * 100)}%"
                                if f["project_status"].get("avg_prob_on_time") is not None
                                else "n/a"
                            ),
                        ]
                        for f in floors
                    ],
                )

        elif key == "cash_flow":
            cf = data.get("cf", {})
            if not cf.get("has_data"):
                doc.add_paragraph("Data not available.")
            else:
                m = cf.get("metrics", {})
                doc.add_paragraph(
                    f"BAC: {_fmt_num(m.get('bac'))}  |  "
                    f"Spent: {_fmt_num(m.get('ac'))} ({m.get('spent_pct', '—')}%)  |  "
                    f"Remaining: {_fmt_num(m.get('remaining'))}  |  "
                    f"Peak month: {m.get('peak_month', '—')} ({_fmt_num(m.get('peak_value'))})"
                )
                months = cf.get("months", [])[:36]
                _docx_table(
                    doc,
                    ["Month", "Planned", "Actual", "Forecast", "Cum. Planned", "Cum. Progress"],
                    [
                        [
                            mo.get("month", ""),
                            _fmt_num(mo.get("planned")),
                            _fmt_num(mo.get("actual")) if mo.get("actual") else "—",
                            _fmt_num(mo.get("forecast")) if mo.get("forecast") else "—",
                            _fmt_num(mo.get("cumulative_planned")),
                            _fmt_num(mo.get("cumulative_progress")),
                        ]
                        for mo in months
                    ],
                )

    # Final footer paragraph
    footer = doc.add_paragraph()
    run = footer.add_run(
        f"\nGenerated by Castor · {data['_project_name']} · {rd} · All figures from live schedule data."
    )
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Public API ────────────────────────────────────────────────────────────────


def generate_report(
    project_id: str,
    fmt: str = "pdf",
    report_type: str = "executive",
    sections: list[str] | None = None,
) -> tuple[bytes, str]:
    """Generate a report and return (bytes, filename).

    Parameters
    ----------
    project_id  : project UUID string
    fmt         : 'pdf' or 'docx'
    report_type : key in REPORT_TYPES (currently 'executive' only)
    sections    : list of section keys to include; None = all for this report type
    """
    from environments.models import Project

    canonical = REPORT_TYPES.get(report_type, ALL_SECTIONS)
    if sections is not None:
        ordered = tuple(k for k in canonical if k in sections)
    else:
        ordered = canonical

    data = gather_report_data(project_id, ordered)

    project = Project.objects.filter(pk=project_id).first()
    safe_name = (project.name if project else "Project").replace(" ", "_")
    today_str = date.today().strftime("%Y-%m-%d")

    if fmt == "docx":
        content = render_docx(data, ordered)
        filename = f"{safe_name}_Executive_Report_{today_str}.docx"
    else:
        content = render_pdf(data, ordered)
        filename = f"{safe_name}_Executive_Report_{today_str}.pdf"

    logger.info(
        "Report generated: project=%s type=%s fmt=%s sections=%d size=%d",
        project_id,
        report_type,
        fmt,
        len(ordered),
        len(content),
    )
    return content, filename
